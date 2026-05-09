import logging
from io import BytesIO

from docx import Document
from docx.oxml import OxmlElement
from docx.shared import Pt

from ..contracts import MeetingAnalysis

logger = logging.getLogger(__name__)


class ExportError(Exception):
    def __init__(self, code: str, message: str):
        self.code = code
        self.message = message
        super().__init__(f"{code}: {message}")


# ─────────────────────────────────────────────────────────────────────────────
# RTL Hebrew helpers — python-docx has no native RTL API.
# We drop into the OOXML layer and append a w:bidi element to the paragraph.
# ─────────────────────────────────────────────────────────────────────────────


def _set_paragraph_rtl(paragraph) -> None:
    pPr = paragraph._p.get_or_add_pPr()
    bidi = OxmlElement("w:bidi")
    pPr.append(bidi)


def _add_rtl_heading(doc: Document, text: str, level: int):
    h = doc.add_heading(text, level=level)
    _set_paragraph_rtl(h)
    return h


def _add_rtl_paragraph(doc: Document, text: str, *, italic: bool = False, size: int = 11):
    p = doc.add_paragraph()
    _set_paragraph_rtl(p)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.complex_script = True
    run.italic = italic
    return p


# ─────────────────────────────────────────────────────────────────────────────


def to_docx(analysis: MeetingAnalysis) -> bytes:
    try:
        doc = Document()

        _add_rtl_heading(doc, "סיכום ישיבה", level=0)

        # Summary
        _add_rtl_heading(doc, "סיכום", level=1)
        _add_rtl_paragraph(doc, analysis.summary)

        # Participants
        _add_rtl_heading(doc, "משתתפים", level=1)
        if analysis.participants:
            for participant in analysis.participants:
                _add_rtl_paragraph(doc, f"• {participant}")
        else:
            _add_rtl_paragraph(doc, "—")

        # Decisions
        _add_rtl_heading(doc, "החלטות", level=1)
        if analysis.decisions:
            for decision in analysis.decisions:
                _add_rtl_paragraph(doc, f"• {decision.text}")
                if decision.context:
                    _add_rtl_paragraph(doc, decision.context, italic=True, size=10)
        else:
            _add_rtl_paragraph(doc, "—")

        # Action items
        _add_rtl_heading(doc, "משימות לביצוע", level=1)
        if analysis.action_items:
            for item in analysis.action_items:
                _add_rtl_paragraph(
                    doc,
                    f"• {item.who} — {item.what} ({item.when})",
                )
        else:
            _add_rtl_paragraph(doc, "—")

        # Open items
        _add_rtl_heading(doc, "נושאים פתוחים", level=1)
        if analysis.open_items:
            for open_item in analysis.open_items:
                _add_rtl_paragraph(doc, f"• {open_item.text}")
        else:
            _add_rtl_paragraph(doc, "—")

        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        result = buffer.read()

        logger.info("docx_export bytes=%d sections=5", len(result))
        return result

    except Exception as exc:
        logger.exception("docx_export_error")
        raise ExportError(
            "EXPORT_FAILED",
            "שגיאה ביצירת המסמך — נסה שוב",
        ) from exc
